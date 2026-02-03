"""Convert .docx files to .txt files without modifying originals."""
from pathlib import Path


def convert_docx_to_txt(docx_path, txt_path):
    """Convert a single .docx file to .txt."""
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed. Installing...")
        import subprocess

        subprocess.check_call(["pip", "install", "python-docx"])
        from docx import Document

    doc = Document(docx_path)

    # Extract all text from paragraphs
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    # Write to txt file
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(full_text))

    return len(full_text)


def main():
    source_dir = Path(r"C:\Users\asafi\Documents\ProjectTT\Plans&Status")

    # Find all .docx files
    docx_files = list(source_dir.glob("*.docx"))

    print(f"Found {len(docx_files)} .docx files to convert")

    for docx_file in docx_files:
        txt_file = docx_file.with_suffix(".txt")
        print(f"Converting: {docx_file.name} -> {txt_file.name}")

        try:
            num_paragraphs = convert_docx_to_txt(docx_file, txt_file)
            print(f"  OK Converted ({num_paragraphs} text blocks)")
        except Exception as e:
            print(f"  ERROR: {e}")

    print("\nConversion complete!")


if __name__ == "__main__":
    main()
